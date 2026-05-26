from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x6B, 0x2D, 0x3E)
DARK   = RGBColor(0x1A, 0x17, 0x14)
MID    = RGBColor(0x5A, 0x54, 0x50)

def set_font(run, size=10.5, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '6B2D3E')
    pBdr.append(bottom)
    pPr.append(pBdr)

lm = Document()

for section in lm.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# NOM
p = lm.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("CÉLIA CARON")
r.font.name = "Calibri"; r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = DARK

# Contact
p = lm.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("celiacaron.crea@gmail.com  |  06 79 68 13 26  |  Issy-les-Moulineaux (92)  |  Portfolio : [URL]")
set_font(r, size=9.5, color=MID)

add_rule(lm)

# Date
p = lm.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Issy-les-Moulineaux, le 22 mai 2026")
set_font(r, size=10, color=MID)

# Destinataire
p = lm.add_paragraph()
p.paragraph_format.space_after = Pt(12)
r = p.add_run("À l'attention du service RH — Marionnaud")
set_font(r, size=10, color=DARK)

# Objet
p = lm.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Objet : Candidature — CDD Graphiste H/F")
set_font(r, size=11, bold=True, color=ACCENT)

# Corps
body = [
    ("normal",
     "Marionnaud est une référence beauté que je suis depuis longtemps — ses campagnes ont une identité visuelle forte, soignée, immédiatement reconnaissable. Intégrer l'équipe design pour contribuer à cette production, c'est exactement le poste que je cherche."),

    ("normal",
     "Pendant 2 ans et demi chez Si Si La Paillette (stage puis alternance), j'ai produit l'intégralité des assets digitaux d'une marque beauté : emailings, newsletters, bannières web, visuels réseaux sociaux, animations motion, packaging. Chaque campagne impliquait de décliner une identité visuelle en dizaines de formats, dans les délais, en suivant une charte précise. J'ai aussi piloté des shootings photo et vidéo de A à Z — du moodboard au montage final. Depuis octobre 2024, mon activité freelance a renforcé ma capacité à travailler en complète autonomie tout en livrant au niveau attendu."),

    ("normal",
     "Ce qui me correspond dans ce poste, c'est la dimension « mi créatif, mi technicien » que vous décrivez — c'est exactement ma façon de fonctionner. J'ai intégré les outils d'IA générative dans ma production courante (Adobe Firefly, Midjourney, Claude, Gemini, Dall-E) pour accélérer les phases d'idéation et d'exploration visuelle. Et si j'ai commencé par une Licence de Physique/Chimie avant de basculer dans le design, c'est parce que j'aime comprendre les systèmes — une rigueur que je mets aujourd'hui au service de l'image."),

    ("normal",
     "L'univers beauté est un fil rouge de mon parcours. Je sais ce qu'il exige visuellement : une image qui attire, qui inspire confiance, qui donne envie. Mon portfolio en témoigne, et je serai ravie d'en discuter avec vous."),
]

for style, text in body:
    p = lm.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_font(r, size=10.5, color=DARK)

# Formule
p = lm.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(20)
r = p.add_run("Cordialement,")
set_font(r, size=10.5, color=DARK)

# Signature
p = lm.add_paragraph()
r = p.add_run("Célia Caron")
set_font(r, size=11, bold=True, color=DARK)

p = lm.add_paragraph()
p.paragraph_format.space_before = Pt(2)
r = p.add_run("celiacaron.crea@gmail.com  |  06 79 68 13 26  |  [Portfolio URL]")
set_font(r, size=9.5, color=MID)

lm.save("/Users/nicolascatteau/Desktop/Célia/job dating/claude portfolio/LM_Celia_Caron_Marionnaud.docx")
print("LM Marionnaud sauvegardée.")
